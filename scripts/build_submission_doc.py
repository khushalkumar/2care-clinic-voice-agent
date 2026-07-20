from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "2care-submission-writeup.docx"
IMG = ROOT / "docs" / "architecture-flowchart.png"

BLUE = "2E74B5"
DARK = "1F4D78"
INK = "0B2545"
MUTED = "555555"
LIGHT = "F2F4F7"
PALE = "E8EEF5"
RED = "9B1C1C"
GREEN = "1F5B3A"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, bold=False, color=None, size=9.5, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.05
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.font.name = "Calibri"
    r._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    r._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    r.font.size = Pt(size)
    r.bold = bold
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell_margins(cell)


def set_table_widths(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")


def add_table(doc, headers, rows, widths, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_widths(table, widths)
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True, color=INK, size=font_size)
        shade(table.rows[0].cells[i], LIGHT)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value), size=font_size)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.08
    p.add_run(text)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.08
    p.add_run(text)
    return p


def add_note(doc, label, text, fill=PALE, color=INK):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_widths(table, [9360])
    cell = table.cell(0, 0)
    shade(cell, fill)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.08
    r = p.add_run(f"{label}: ")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(color)
    r2 = p.add_run(text)
    r2.font.color.rgb = RGBColor.from_string(color)
    cell_margins(cell, 140, 160, 140, 160)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def set_run(run, size=11, color=None, bold=None, italic=None, font="Calibri"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_para(doc, text="", after=6, before=0, color=None, size=11, bold=False, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    set_run(p.add_run(text), size=size, color=color, bold=bold, italic=italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def setup_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for name in ("List Bullet", "List Bullet 2", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.line_spacing = 1.08


def add_header_footer(section):
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    set_run(header.add_run("2care.ai | Voice AI Engineer Assignment"), size=9, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(footer.add_run("2care Clinic Voice Agent | Submission write-up"), size=8.5, color=MUTED)


def add_architecture_image():
    from PIL import Image, ImageDraw, ImageFont

    width, height = 1800, 650
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    try:
        font = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial.ttf", 23)
        small = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial.ttf", 18)
    except OSError:
        font = ImageFont.load_default()
        small = font
    boxes = [
        (40, 250, 230, 360, "Caller\nEnglish / Hindi / Hinglish", "E8EEF5"),
        (300, 250, 500, 360, "Retell\nvoice agent", "DCEAF7"),
        (570, 250, 770, 360, "AWS ALB\nWAF rate limit", "E8EEF5"),
        (840, 180, 1090, 430, "FastAPI on\nECS Fargate\n\nHMAC + replay\n+ tool contracts", "DCEAF7"),
        (1160, 80, 1400, 205, "RDS PostgreSQL\nstate + reservations", "EAF3EE"),
        (1160, 260, 1400, 375, "Cliniko PMS\nlive availability", "FFF4D6"),
        (1160, 440, 1400, 555, "SQS + DLQ\nreconciliation", "F6E9E9"),
        (1470, 440, 1740, 555, "Worker\nwrite recovery", "F6E9E9"),
    ]
    for x1, y1, x2, y2, label, fill in boxes:
        draw.rounded_rectangle((x1, y1, x2, y2), radius=18, fill=f"#{fill}", outline="#1F4D78", width=3)
        lines = label.split("\n")
        total = len(lines) * 28
        y = (y1 + y2 - total) / 2
        for line in lines:
            bbox = draw.textbbox((0, 0), line, font=font)
            draw.text(((x1+x2-bbox[2])/2, y), line, fill="#0B2545", font=font)
            y += 28
    arrows = [((230, 305), (300, 305)), ((500, 305), (570, 305)), ((770, 305), (840, 305)),
              ((1090, 230), (1160, 145)), ((1090, 305), (1160, 315)), ((1090, 360), (1160, 500)),
              ((1400, 500), (1470, 500))]
    for start, end in arrows:
        draw.line((start, end), fill="#2E74B5", width=5)
        ex, ey = end
        draw.polygon([(ex, ey), (ex-14, ey-8), (ex-14, ey+8)], fill="#2E74B5")
    draw.text((40, 35), "Production-oriented architecture", fill="#0B2545", font=font)
    draw.text((40, 90), "Cliniko owns clinic records; PostgreSQL owns state, conflict protection, and recovery.", fill="#555555", font=small)
    image.save(IMG)


def main():
    add_architecture_image()
    doc = Document()
    setup_styles(doc)
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    add_header_footer(section)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    set_run(p.add_run("2care.ai"), size=12, color=MUTED, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    set_run(p.add_run("Production Voice AI Receptionist"), size=25, color=INK, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    set_run(p.add_run("Voice AI Engineer Assignment | Submission Write-up"), size=14, color=MUTED)
    add_note(doc, "Submission status", "The implementation is live and independently callable. This document reports the completed engineering work and the observed pilot evidence honestly. The complete 17-scenario live bake-off required by the assignment has not yet been run; the limitations section identifies this explicitly.")

    add_heading(doc, "1. Executive Summary", 1)
    add_para(doc, "This project implements a production-oriented bilingual voice receptionist for two Physiotattva demonstration clinics in Bengaluru: Jayanagar and Indiranagar. Patients can book, reschedule, cancel, recover from dropped calls, request human follow-up, and use English, Hindi, or natural Hinglish code-switching.")
    add_para(doc, "The live path is Retell -> Twilio Elastic SIP Trunk -> AWS ALB -> FastAPI on ECS Fargate -> RDS PostgreSQL and Cliniko. The service is not a single happy-path webhook: it uses persisted call state, opaque availability tokens, write-time conflict protection, idempotency, replay-safe authentication, follow-up logging, and reconciliation for uncertain PMS writes.")
    add_para(doc, "The repository, deployment configuration, prompt, tool contracts, test suite, evaluation corpus, runbooks, and decision register are available in the GitHub repository:")
    add_para(doc, "https://github.com/khushalkumar/2care-clinic-voice-agent", color=BLUE)

    add_heading(doc, "2. Live Test Details", 1)
    add_table(doc, ["Item", "Value"], [
        ("Live phone number", "+1 417 742 8846"),
        ("Voice platform", "Retell"),
        ("Telephony", "Twilio Elastic SIP Trunk"),
        ("Backend", "AWS ECS Fargate behind an ALB"),
        ("Database", "AWS RDS PostgreSQL"),
        ("PMS", "Cliniko, shard au1"),
        ("Test browser", "https://khushalkumar.github.io/2care-clinic-voice-agent/"),
        ("Staging deployment", "GitHub Actions deploy-staging run #32, successful"),
    ], [2600, 6760])
    add_note(doc, "Live test guidance", "Ask for an appointment in English, Hindi, or Hinglish. Give a full name before booking. Try a named branch, a natural time such as 'Thursday morning' or 'after work around 4:30', and ask for a human if needed. The browser demo requires microphone permission. Never share generated web-call URLs because they contain short-lived access tokens.", fill="FFF4D6", color="7A5A00")

    add_heading(doc, "3. What Was Built", 1)
    add_bullet(doc, "Full appointment lifecycle: booking, rescheduling, cancellation, conflict handling, and human follow-up.")
    add_bullet(doc, "Two sourced clinic branches with branch-specific appointment types and practitioners in Cliniko.")
    add_bullet(doc, "Live availability queries with fresh searches whenever branch, practitioner, date, service, or time constraints change.")
    add_bullet(doc, "Opaque, short-lived availability tokens. Mutation tools do not accept caller-controlled raw slot timestamps.")
    add_bullet(doc, "PostgreSQL-backed call sessions, checkpoints, idempotency, replay prevention, local reservations, and recovery state.")
    add_bullet(doc, "Explicit full-name authorization, including shared-phone disambiguation and new-patient booking.")
    add_bullet(doc, "Dropped-call continuation and missed-callback context through persisted session state.")
    add_bullet(doc, "Timeout-after-write reconciliation so the agent does not claim failure when the PMS write may have succeeded, and does not claim success without a definitive PMS record.")
    add_bullet(doc, "Request hardening: HMAC authentication, request limits, strict content type, security headers, default-deny CORS, replay protection, and WAF rate limiting in the infrastructure profile.")

    add_heading(doc, "4. Architecture", 1)
    doc.add_picture(str(IMG), width=Inches(6.5))
    cap = doc.add_paragraph("Figure 1. Live request and recovery path.")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(cap.runs[0], size=9, color=MUTED, italic=True)
    add_para(doc, "Cliniko is the source of truth for clinic records, practitioners, appointment types, live availability, and final appointment records. PostgreSQL owns conversation state, reservations, idempotency, auditability, follow-ups, and recovery. A booking is spoken as confirmed only after local conflict protection and a definitive Cliniko write both succeed.")

    add_heading(doc, "5. Clinic and PMS", 1)
    add_para(doc, "The demonstration clinic is Physiotattva, selected because its official pages provide two Bengaluru branches, named practitioners, services, addresses, and a plausible physiotherapy scheduling model. The implementation keeps the Cliniko subset intentionally small so the behavior is testable and understandable.")
    add_table(doc, ["Branch", "Address used in the demo", "Purpose"], [
        ("Jayanagar", "75, 8th Main Road, Jaya Nagar 1st Block, Bengaluru 560011", "Branch-specific booking and triage"),
        ("Indiranagar", "1st floor, 3478/A, 14th Main Road, HAL 2nd Stage, Bengaluru 560038", "Cross-branch earliest-slot search"),
    ], [1500, 4700, 3160])
    add_para(doc, "Public source pages are recorded in the repository decision register. Exact practitioner rotas and operational policies that were not publicly verifiable are treated as demo operational data, not represented as official clinic policy.")

    add_heading(doc, "6. Platform Choice", 1)
    add_para(doc, "The assignment's explicit platform gate asks us to pick Retell or Bolna and build one platform fully live. Vapi and LiveKit were still compared as engineering alternatives, but they were not selected as the final assignment platform under that narrower requirement.")
    add_table(doc, ["Dimension", "Retell", "Bolna", "Vapi", "LiveKit"], [
        ("Assignment fit", "Pass", "Pass", "Narrow-clause fail", "Narrow-clause fail"),
        ("Hindi / Hinglish", "Multilingual agent locales and Hindi support", "Strong Indian-provider options", "Provider-dependent", "Provider-dependent"),
        ("Tool calling", "Hosted schemas, timeouts, retries, simulation", "Custom function and graph-agent support", "Flexible server tools", "Code-native tools"),
        ("Latency evidence", "Component latency and percentiles exposed", "Trace timestamps available", "Provider-dependent", "Deep code-level control"),
        ("Telephony", "Managed numbers and custom SIP", "Twilio, Plivo, Exotel, Vobiz", "Built-in and imported numbers", "SIP and carrier integrations"),
        ("Three-day delivery risk", "Low", "Low-moderate", "Low", "High"),
    ], [1500, 2200, 2200, 1730, 1730], font_size=8.6)
    add_para(doc, "Retell was chosen because it offered the best combination of assignment compliance, managed telephony, structured custom-function schemas, interruption controls, multilingual configuration, and an observable testing surface within the deadline. The decision is not based on a claim that Retell is universally better. The repository includes a comparison and a re-benchmark plan; actual live evidence should decide a production platform at larger scale.")

    add_heading(doc, "7. Conversation and Prompt Logic", 1)
    add_para(doc, "The complete versioned prompt is in integrations/voice/retell/prompt.md. Its key rules are:")
    add_number(doc, "Call clinic_catalog once and use only returned identifiers; never invent practitioner, service, branch, or appointment-type IDs.")
    add_number(doc, "Collect the caller phone number, bootstrap a persisted call session, and use the returned session_id for all stateful tools.")
    add_number(doc, "Ask for and confirm the full name even when the phone number is recognized. Do not expose appointment details when lookup is ambiguous.")
    add_number(doc, "Search live availability before offering times and re-search after every changed constraint. Offer at most three backend-generated spoken labels.")
    add_number(doc, "Repeat branch, practitioner, and local India time before mutation. Confirm success only when the backend returns confirmed.")
    add_number(doc, "For rescheduling, identify the appointment, search fresh availability, and use a fresh opaque availability token. For cancellation, repeat appointment details and obtain confirmation.")
    add_number(doc, "Mirror pure English, pure Hindi, and natural Hinglish. Do not drift languages without a caller cue, and do not translate names, branch names, dates, or times.")
    add_number(doc, "Use one concise holding phrase while a tool runs. Log human follow-up for clinical, unsupported, or human-requested cases without pretending an immediate transfer occurred.")

    add_heading(doc, "8. Evaluation Harness and Current Evidence", 1)
    add_para(doc, "The repository contains 17 versioned multi-turn scenarios across English, Hindi, and Hinglish, a redaction validator, a renderer, and per-language reporting. The strict validator is designed to prevent a partial or blended result from being presented as complete coverage.")
    add_para(doc, "Existing live evidence was evaluated from eight completed Retell phone calls made on 2026-07-20. One web call where the caller never joined was excluded. The calls were exploratory rather than pre-assigned one-to-one to the 17 scenario IDs.")
    add_table(doc, ["Observed bucket", "Calls", "Definitive bookings", "ASR p50", "LLM p50", "TTS p50", "E2E p50"], [
        ("English", "6", "0", "195 ms", "1,023 ms", "384 ms", "2,265 ms"),
        ("Hindi / Hinglish mixed", "2", "0", "218 ms", "1,912 ms", "483 ms", "3,424 ms"),
    ], [2300, 800, 1400, 1250, 1250, 1250, 1110], font_size=8.8)
    add_para(doc, "Network latency was unavailable in the exported Retell traces and is reported as unavailable rather than inferred. The complete redacted pilot report is available at:")
    add_para(doc, "https://github.com/khushalkumar/2care-clinic-voice-agent/blob/main/evals/reports/observed-pilot-report.md", color=BLUE)
    add_note(doc, "Interpretation", "The sample demonstrates real calls and real failure handling, but it does not prove complete assignment compliance. In particular, it contains no definitive booking, does not cover every required scenario, and does not validate the full English/Hindi/Hinglish bake-off. These are live-evidence gaps, not hidden claims.", fill="F6E9E9", color=RED)

    add_heading(doc, "9. Reproducibility and Deployment", 1)
    add_para(doc, "A clean clone can run the test suite with Python 3.11, PostgreSQL 16 test binaries, and the locked requirements. The repository uses venv for local setup and includes CI for tests, lint, type checks, infrastructure validation, secret scanning, and container scanning.")
    add_para(doc, "Local setup:")
    for line in [
        "python3.11 -m venv venv",
        "source venv/bin/activate",
        "python -m pip install -r requirements.lock",
        "python -m pip install --no-deps -e .",
        "pytest",
        "ruff check app tests migrations",
        "mypy app",
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.0
        set_run(p.add_run(line), size=9.5, color=INK, font="Courier New")
    add_para(doc, "Deployment is defined in infra/terraform and .github/workflows. Staging uses one ECS task, private RDS, ALB access, Secrets Manager, KMS, SQS/DLQ, CloudWatch, and CloudTrail resources. Production profile settings enable stronger isolation, multiple tasks, Multi-AZ RDS, backups, and deletion protection.")

    add_heading(doc, "10. Known Limitations and Next Verification", 1)
    add_bullet(doc, "Run every versioned scenario in the required language mode and record redacted scenario-level measurements.")
    add_bullet(doc, "Run a final live synthetic booking, reschedule, cancellation, conflict, and human-follow-up canary after the current deployment.")
    add_bullet(doc, "Capture controlled network timing and backend spans so network latency can be reported separately from platform component latency.")
    add_bullet(doc, "Repeat the live calls after the latest patient-booking and language handling fixes; the observed pilot includes earlier failures and should not be treated as a current success rate.")
    add_bullet(doc, "ACM/HTTPS with a controlled domain remains a production hardening step. The assignment staging demo currently uses an AWS ALB endpoint without purchasing a domain.")
    add_para(doc, "These limitations are the reason the README does not claim that the full DOCX acceptance gate has been completed.")

    add_heading(doc, "11. Submission Email Draft", 1)
    add_para(doc, "To: tech@2care.ai, p@2care.ai, s@2care.ai")
    add_para(doc, "Subject: Voice AI Engineer Assignment - 2care Clinic Voice Agent", bold=True)
    add_para(doc, "Hello 2care.ai team,\n\nI have completed the implementation portion of the Voice AI Engineer assignment and built a live bilingual clinic receptionist using Retell, Twilio, AWS, PostgreSQL, and Cliniko. The write-up, architecture, prompt logic, deployment instructions, tests, and evaluation materials are available in the links below.")
    add_para(doc, "GitHub repository: https://github.com/khushalkumar/2care-clinic-voice-agent\nWrite-up: https://github.com/khushalkumar/2care-clinic-voice-agent/blob/main/README.md\nPrompt: https://github.com/khushalkumar/2care-clinic-voice-agent/blob/main/integrations/voice/retell/prompt.md\nLive test number: +1 417 742 8846\nBrowser test: https://khushalkumar.github.io/2care-clinic-voice-agent/")
    add_para(doc, "The repository also includes an observed pilot report from eight completed calls. It is clearly labeled as exploratory evidence; the complete 17-scenario live evaluation remains identified as pending rather than overstated.")
    add_para(doc, "Regards,\nKhushal Kumar")

    add_heading(doc, "Appendix A. Source and Repository Map", 1)
    add_table(doc, ["Artifact", "Location"], [
        ("Assignment source", "docs/assignment/assignment.txt and docs/assignment/original-assignment.docx"),
        ("Prompt", "integrations/voice/retell/prompt.md"),
        ("Tool contracts", "integrations/voice/tool-contracts.json"),
        ("Architecture", "docs/architecture.md"),
        ("Evaluation protocol", "docs/evaluation.md"),
        ("Scenario corpus", "evals/scenarios/core.json"),
        ("Observed report", "evals/reports/observed-pilot-report.md"),
        ("Security design", "docs/security-and-privacy.md"),
        ("Deployment runbook", "docs/runbooks/deployment.md"),
    ], [2600, 6760])
    add_para(doc, "Prepared from repository state at commit 3a03338 on 2026-07-20.", after=0, color=MUTED, size=9, italic=True)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
